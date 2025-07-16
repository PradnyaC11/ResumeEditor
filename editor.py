import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from datetime import date
from docx2pdf import convert
# import win32com.client
from docx.shared import Pt


def replace_text_in_paragraph(paragraph, key, value):
    inline = paragraph.runs
    for item in inline:
        # print(item.text)
        if key in item.text:
            item.text = item.text.replace(key, value)


output_file_path = "D:/Resumes/Resumes/Tailored/Resume_Pradnya_Chaudhari.docx"
pdf_file_path = "D:/Resumes/Resumes/Tailored/Resume_Pradnya_Chaudhari.pdf"
resume_template = docx.Document("D:/Resumes/ResumeEditor/Resume_Template_8_Java_Angular.docx")
# resume_template = docx.Document("D:/Resumes/ResumeEditor/Resume_Template_8_Java_React.docx")

variables = {
    # "Spring Boot": "Spring-Boot",
    # "React": "ReactJS",
    # "NodeJS": "Node.js",
    # "Angular": "AngularJS",
    # "HTML": "HTML5",
    # "PostgreSQL": "Postgres",
    # "REST": "RESTful"
    # "microservices": "micro-services",
    # "Go": "Golang",
    # "CSS": "CSS3"
    # "D3.js": "D3"
    # "full-stack": "full stack"
    # "CI/CD": "CICD"
    # "Postman": "Postman, IntelliJ"
}


# replace_text_in_word(resume_template, variables, output_file_path)

# replace_text_in_document(resume_template, variables)

for variable_key, variable_value in variables.items():
    for paragraph in resume_template.paragraphs:
        replace_text_in_paragraph(paragraph, variable_key, variable_value)
        # replace_text_preserve_formatting(paragraph, variable_key, variable_value)
# Save the document
resume_template.save(output_file_path)

convert(output_file_path, pdf_file_path)
